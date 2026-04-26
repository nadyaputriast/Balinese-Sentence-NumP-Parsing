import io
import re
import unicodedata
import pandas as pd
import streamlit as st

from core import cyk_algorithm, convert_to_cnf, remove_epsilon_productions, remove_unit_productions
from grammar import RULES_CFG
from utils import stats_manager
from docx import Document


def read_to_dataframe(uploaded_file) -> tuple[pd.DataFrame | None, str | None]:
    """
    Baca berbagai format file → DataFrame dengan kolom 'kalimat'.
    Format yang didukung: CSV, Excel (.xlsx/.xls), Word (.docx), Plain Text (.txt)
    Returns: (df, error_message)
    """
    name = uploaded_file.name.lower()

    try:
        # CSV
        if name.endswith('.csv'):
            raw = uploaded_file.read()
            uploaded_file.seek(0)
            sample = raw[:2048].decode('utf-8', errors='ignore')
            sep = '\t' if sample.count('\t') > sample.count(',') else ','
            df = pd.read_csv(io.BytesIO(raw), sep=sep)

        # Excel
        elif name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)

        # Word (.docx)
        elif name.endswith('.docx'):
            doc = Document(io.BytesIO(uploaded_file.read()))

            # Prioritas 1: tabel yang punya header 'kalimat'
            for table in doc.tables:
                headers = [c.text.strip().lower() for c in table.rows[0].cells]
                if 'kalimat' in headers:
                    rows = [
                        {h: table.rows[r].cells[i].text.strip() for i, h in enumerate(headers)}
                        for r in range(1, len(table.rows))
                    ]
                    df = pd.DataFrame(rows)
                    df = df[df['kalimat'].str.strip().ne('')]
                    return _normalize_columns(df, uploaded_file.name), None

            # Prioritas 2: paragraf & list (bullet/numbered)
            kalimat_list = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Strip prefix list: "1.", "1)", "-", "•", "*"
                text = re.sub(r'^(\d+[\.\)]|[-•*])\s*', '', text).strip()
                if text:
                    kalimat_list.append(text)

            if not kalimat_list:
                return None, f"Tidak ada teks ditemukan di '{uploaded_file.name}'"

            df = pd.DataFrame({'kalimat': kalimat_list})

        # Plain text (.txt)
        elif name.endswith('.txt'):
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            lines = [l.strip() for l in content.splitlines() if l.strip()]
            if not lines:
                return None, f"File '{uploaded_file.name}' kosong atau tidak terbaca"
            df = pd.DataFrame({'kalimat': lines})

        else:
            ext = '.' + name.rsplit('.', 1)[-1] if '.' in name else 'unknown'
            return None, (
                f"Format `{ext}` belum didukung. "
                f"Format yang bisa: `.csv`, `.xlsx`, `.xls`, `.docx`, `.txt`"
            )

    except Exception as e:
        return None, f"Gagal membaca '{uploaded_file.name}': {str(e)}"

    df.columns = df.columns.str.strip()
    return _normalize_columns(df, uploaded_file.name), None


def _normalize_columns(df: pd.DataFrame, filename: str) -> pd.DataFrame:
    """
    Pastikan kolom 'kalimat' ada.
    Kalau tidak ada, gunakan kolom pertama sebagai fallback + warning.
    """
    if 'kalimat' not in df.columns:
        first_col = df.columns[0]
        df = df.rename(columns={first_col: 'kalimat'})
        st.warning(
            f"⚠️ Kolom 'kalimat' tidak ditemukan di **{filename}**. "
            f"Menggunakan kolom pertama: **'{first_col}'**"
        )
    return df


def process_files(
    uploaded_files,
    kamus_dasar,
    stemmer_func
) -> tuple[pd.DataFrame | None, str | None]:
    """
    Proses satu atau beberapa file sekaligus.
    - Semua kolom original dipertahankan
    - Tambah kolom 'sumber' (nama file) di awal
    - Tambah kolom 'status' (VALID/INVALID) di akhir
    Returns: (df_gabungan, error_message)
    """
    # Siapkan grammar sekali untuk semua file
    cfg_cleaned = remove_epsilon_productions(RULES_CFG)
    cfg_cleaned = remove_unit_productions(cfg_cleaned)
    cnf_grammar = convert_to_cnf(cfg_cleaned)

    all_dfs = []

    for uploaded_file in uploaded_files:
        st.caption(f"⏳ Membaca **{uploaded_file.name}**...")

        df, err = read_to_dataframe(uploaded_file)
        if err:
            st.error(f"❌ {uploaded_file.name}: {err}")
            continue

        total = len(df)
        if total == 0:
            st.warning(f"⚠️ **{uploaded_file.name}** tidak memiliki data, dilewati.")
            continue

        st.caption(f"✅ **{uploaded_file.name}** — {total} baris ditemukan")

        results = []
        progress_bar = st.progress(0, text=f"Memproses {uploaded_file.name}...")

        for i, row in df.iterrows():
            sentence_raw = str(row['kalimat']).lower().strip()
            sentence_normalized = (
                unicodedata.normalize('NFKD', sentence_raw)
                .encode('ASCII', 'ignore')
                .decode('utf-8')
            )
            sentence_final, _ = stemmer_func(sentence_normalized, kamus_dasar)
            words = sentence_final.split()

            if not words:
                results.append("INVALID")
            else:
                is_valid, _, _ = cyk_algorithm(cnf_grammar, words)
                results.append("VALID" if is_valid else "INVALID")

            progress_bar.progress((i + 1) / total, text=f"Memproses {uploaded_file.name}...")

        # Kolom: sumber | kolom original | status
        if len(uploaded_files) > 1:
            df.insert(0, 'sumber', uploaded_file.name)
        df['status'] = results

        all_dfs.append(df)

    if not all_dfs:
        return None, "Tidak ada file yang berhasil diproses."

    return pd.concat(all_dfs, ignore_index=True), None


def to_excel_bytes(df: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()

    writer = pd.ExcelWriter(buffer, engine='openpyxl')
    df.to_excel(writer, index=False, sheet_name="Hasil Validasi")

    worksheet = writer.sheets["Hasil Validasi"]

    from openpyxl.styles import PatternFill, Font
    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    red_fill   = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    green_font = Font(color="276221")
    red_font   = Font(color="9C0006")

    headers = [cell.value for cell in worksheet[1]]
    if 'status' in headers:
        status_col = headers.index('status') + 1
        for row in worksheet.iter_rows(min_row=2, min_col=status_col, max_col=status_col):
            for cell in row:
                if cell.value == "VALID":
                    cell.fill = green_fill
                    cell.font = green_font
                elif cell.value == "INVALID":
                    cell.fill = red_fill
                    cell.font = red_font

    for col in worksheet.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=0)
        worksheet.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)

    writer.close()
    return buffer.getvalue()